"""
Genera contrato-servicios.docx y contrato-servicios.pdf
con formato profesional.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── helpers ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    return h


def add_paragraph_justified(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    return p


def add_checkbox_line(doc, text, checked=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    mark = '\u2611' if checked else '\u2610'
    run = p.add_run(f'{mark}  {text}')
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    return p


def add_blank_field(doc, label, width_chars=40):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{label}: ')
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.bold = True
    run2 = p.add_run('_' * width_chars)
    run2.font.size = Pt(10.5)
    run2.font.name = 'Calibri'
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
    return p


# ── build DOCX ───────────────────────────────────────────

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # ── Title ──
    title = doc.add_heading('CONTRATO DE LOCACI\u00d3N DE SERVICIOS\nDE DESARROLLO WEB Y DATOS', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
        run.font.size = Pt(18)

    # ── Separator ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2501' * 60)
    run.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6)
    run.font.size = Pt(8)

    # ── Parties ──
    add_heading_styled(doc, 'Partes', level=2)

    add_paragraph_justified(doc,
        'EL PRESTADOR: Juan Ignacio Barranco, DNI N.\u00b0 ____________, '
        'con domicilio en Maip\u00fa, Mendoza, correo electr\u00f3nico jbarranco.sistemas@gmail.com, '
        'tel\u00e9fono +54 9 261 642-3730, en adelante \u201cEL DESARROLLADOR\u201d.')

    add_paragraph_justified(doc,
        'EL COMITENTE: ____________________________________________, '
        'DNI/CUIT N.\u00b0 ____________, con domicilio comercial en '
        '____________________________________________, correo electr\u00f3nico '
        '________________________, tel\u00e9fono ________________________, en adelante \u201cEL CLIENTE\u201d.')

    add_paragraph_justified(doc,
        'En la ciudad de Mendoza, a los ______ d\u00edas del mes de ______________ de 20____, '
        'las partes acuerdan celebrar el presente contrato de locaci\u00f3n de servicios conforme '
        'a los art\u00edculos 1251 a 1261 del C\u00f3digo Civil y Comercial de la Naci\u00f3n, sujeto a las '
        'siguientes cl\u00e1usulas:')

    # ── CL\u00c1USULA 1 ──
    add_heading_styled(doc, 'Cl\u00e1usula Primera \u2014 Objeto', level=2)
    add_paragraph_justified(doc,
        'El Desarrollador se compromete a dise\u00f1ar, desarrollar y entregar al Cliente un sitio '
        'web profesional y/o servicios de datos seg\u00fan las especificaciones detalladas en el '
        'Anexo A \u2014 Alcance del proyecto, que forma parte integrante de este contrato.')
    p = add_paragraph_justified(doc, 'El plan/servicio contratado es:')
    add_checkbox_line(doc, 'Starter ($79.200 \u2014 lanzamiento)')
    add_checkbox_line(doc, 'Profesional ($224.000 \u2014 lanzamiento)')
    add_checkbox_line(doc, 'Premium ($384.000 \u2014 lanzamiento)')
    add_checkbox_line(doc, 'Analytics Setup ($44.000)')
    add_checkbox_line(doc, 'Data Pro ($152.000)')
    add_checkbox_line(doc, 'Data Enterprise (a medida)')

    # ── CL\u00c1USULA 2 ──
    add_heading_styled(doc, 'Cl\u00e1usula Segunda \u2014 Alcance y exclusiones', level=2)
    add_paragraph_justified(doc, '2.1. Incluye:', bold=True)
    add_paragraph_justified(doc,
        'El trabajo comprende exclusivamente lo detallado en el Anexo A. '
        'A modo de referencia seg\u00fan el plan contratado:')

    add_paragraph_justified(doc, 'Plan Starter ($79.200 \u2014 lanzamiento):', bold=True)
    for item in [
        'Sitio web profesional a medida (hasta 5 secciones)',
        'Dise\u00f1o responsive (celular, PC, tablet)',
        'Formulario de contacto funcional',
        'Bot\u00f3n de WhatsApp',
        'Configuraci\u00f3n de Google My Business',
        'SEO b\u00e1sico (meta tags, estructura, schema local)',
        'Configuraci\u00f3n de Google Analytics 4',
        'Registro de dominio .com.ar (primer a\u00f1o)',
        'Publicaci\u00f3n en hosting sin costo mensual',
    ]:
        add_bullet(doc, item)

    add_paragraph_justified(doc, 'Plan Profesional ($224.000 \u2014 lanzamiento) \u2014 incluye todo lo anterior m\u00e1s:', bold=True)
    for item in [
        'Sitio de 5 o m\u00e1s secciones',
        'Sistema de reservas o turnos online integrado',
        'SEO avanzado + configuraci\u00f3n completa de Google Analytics',
        'Capacitaci\u00f3n b\u00e1sica para gestionar contenido',
    ]:
        add_bullet(doc, item)

    add_paragraph_justified(doc, 'Plan Premium ($384.000 \u2014 lanzamiento) \u2014 incluye todo lo anterior m\u00e1s:', bold=True)
    for item in [
        'Cat\u00e1logo de productos/servicios con integraci\u00f3n de Mercado Pago',
        'Versi\u00f3n biling\u00fce del sitio (espa\u00f1ol e ingl\u00e9s)',
        'Configuraci\u00f3n avanzada de Google Analytics con eventos personalizados',
        'Formularios inteligentes con notificaciones autom\u00e1ticas por email',
    ]:
        add_bullet(doc, item)

    # Data services
    add_paragraph_justified(doc, 'Servicios de datos:', bold=True)

    add_paragraph_justified(doc, 'Analytics Setup ($44.000):', bold=True)
    for item in [
        'Configuraci\u00f3n de Google Analytics 4 con eventos personalizados',
        'Dashboard b\u00e1sico de m\u00e9tricas clave',
        'Integraci\u00f3n con Google Search Console',
        'Informe inicial de rendimiento',
    ]:
        add_bullet(doc, item)

    add_paragraph_justified(doc, 'Data Pro ($152.000):', bold=True)
    for item in [
        'Todo lo de Analytics Setup',
        'Dashboard avanzado con m\u00e9tricas de negocio',
        'Integraci\u00f3n de m\u00faltiples fuentes de datos',
        'Automatizaci\u00f3n de reportes peri\u00f3dicos',
        'Capacitaci\u00f3n para interpretar datos',
    ]:
        add_bullet(doc, item)

    add_paragraph_justified(doc, 'Data Enterprise (desde $304.000 \u2014 a medida):', bold=True)
    for item in [
        'Todo lo de Data Pro',
        'Pipelines de datos personalizados',
        'Data warehouse / data lake seg\u00fan necesidad',
        'Modelos predictivos y an\u00e1lisis avanzado',
        'Soporte continuo y mantenimiento de infraestructura de datos',
    ]:
        add_bullet(doc, item)

    add_paragraph_justified(doc, '2.2. No incluye (salvo acuerdo escrito por separado):', bold=True)
    for item in [
        'Creaci\u00f3n de contenido (textos, fotograf\u00edas, logotipo) \u2014 el Cliente los provee',
        'Redacci\u00f3n o traducci\u00f3n de textos m\u00e1s all\u00e1 de lo indicado en el plan',
        'Gesti\u00f3n de redes sociales',
        'Campa\u00f1as de publicidad (Google Ads, Meta Ads, etc.)',
        'Funcionalidades no especificadas en el Anexo A',
        'Desarrollo de aplicaciones m\u00f3viles nativas',
        'Integraciones con sistemas de terceros no mencionados',
    ]:
        add_bullet(doc, item)

    # ── CL\u00c1USULA 3 ──
    add_heading_styled(doc, 'Cl\u00e1usula Tercera \u2014 Precio y forma de pago', level=2)
    add_blank_field(doc, 'Precio total', 30)
    add_paragraph_justified(doc, 'Forma de pago acordada:')
    add_checkbox_line(doc, 'Pago \u00fanico: 50% al aprobar el prototipo, 50% al recibir el sitio terminado.')
    add_checkbox_line(doc, '3 cuotas de $__________ cada una. Primera al aprobar el prototipo, las siguientes a 30 y 60 d\u00edas.')
    add_checkbox_line(doc, 'Otra: ____________________________________________')
    add_paragraph_justified(doc,
        '3.1. Medios de pago aceptados: transferencia bancaria, efectivo, Mercado Pago.')
    add_paragraph_justified(doc,
        '3.2. Stablecoins aceptadas: USDT, USDC, DAI (red Tron o Polygon).')
    add_paragraph_justified(doc,
        '3.3. Los pagos en criptomonedas se liquidan al tipo de cambio del momento de acreditaci\u00f3n.')
    add_paragraph_justified(doc,
        '3.4. El Desarrollador no emite factura fiscal al momento de la firma del presente contrato. '
        'Si en el futuro se inscribiera como contribuyente, podr\u00e1 emitir el comprobante correspondiente. '
        'Esto no afecta la validez del presente contrato ni las obligaciones asumidas por las partes.')
    add_paragraph_justified(doc,
        '3.5. La falta de pago en los plazos acordados habilita al Desarrollador a suspender el trabajo '
        'en curso y retener la entrega de archivos hasta su regularizaci\u00f3n.')
    add_paragraph_justified(doc,
        '3.6. Los precios indicados corresponden a la promoci\u00f3n de lanzamiento (20% de descuento sobre '
        'precios de lista), v\u00e1lida para los primeros 10 clientes. Precios de lista: Starter $99.000, '
        'Profesional $280.000, Premium $480.000.')

    # ── CL\u00c1USULA 4 ──
    add_heading_styled(doc, 'Cl\u00e1usula Cuarta \u2014 Plazos', level=2)
    add_paragraph_justified(doc,
        'Prototipo visual: El Desarrollador entregar\u00e1 un prototipo visual navegable dentro de las '
        '48 horas h\u00e1biles siguientes a la recepci\u00f3n del contenido completo por parte del Cliente.')
    add_paragraph_justified(doc, 'Sitio terminado (desde aprobaci\u00f3n del prototipo):')
    add_bullet(doc, 'Plan Starter: 7 (siete) d\u00edas h\u00e1biles')
    add_bullet(doc, 'Plan Profesional: 14 (catorce) d\u00edas h\u00e1biles')
    add_bullet(doc, 'Plan Premium: 21 (veinti\u00fan) d\u00edas h\u00e1biles')
    add_paragraph_justified(doc,
        'Los plazos se computan desde que el Cliente provee la totalidad del contenido y aprueba '
        'el prototipo. Demoras atribuibles al Cliente (falta de respuesta, contenido incompleto, '
        'cambios de alcance) suspenden el c\u00f3mputo del plazo.')
    add_paragraph_justified(doc,
        'El Cliente dispone de 2 (dos) rondas de revisi\u00f3n incluidas en el precio. Cambios de '
        'estructura, funcionalidad o alcance se consideran trabajo adicional.')

    # ── CL\u00c1USULA 5 ──
    add_heading_styled(doc, 'Cl\u00e1usula Quinta \u2014 Obligaciones del Cliente', level=2)
    add_paragraph_justified(doc, 'El Cliente se compromete a:')
    add_bullet(doc, 'Proveer todo el contenido necesario dentro de los 5 d\u00edas h\u00e1biles posteriores a la firma.')
    add_bullet(doc, 'Designar una persona de contacto con capacidad de decisi\u00f3n.')
    add_bullet(doc, 'Responder consultas del Desarrollador en un plazo no mayor a 3 d\u00edas h\u00e1biles.')
    add_bullet(doc, 'Facilitar acceso a cuentas necesarias (Google My Business, dominio, Mercado Pago, etc.).')
    add_bullet(doc, 'Abonar el precio en los plazos convenidos.')

    # ── CL\u00c1USULA 6 ──
    add_heading_styled(doc, 'Cl\u00e1usula Sexta \u2014 Propiedad intelectual', level=2)
    add_paragraph_justified(doc,
        'Una vez cancelado el 100% del precio, la totalidad del c\u00f3digo fuente, dise\u00f1o, archivos '
        'y contenido desarrollado pasan a ser propiedad exclusiva del Cliente. El Cliente puede '
        'usarlo, modificarlo, transferirlo o contratar a otro desarrollador sin restricci\u00f3n alguna.')
    add_paragraph_justified(doc,
        'El Desarrollador se reserva el derecho de utilizar capturas del sitio en su portfolio '
        'profesional, salvo prohibici\u00f3n expresa por escrito del Cliente.')
    add_paragraph_justified(doc,
        'Las im\u00e1genes de stock utilizadas (Unsplash, Pexels, etc.) se rigen por sus propias '
        'licencias de uso libre. El Desarrollador informar\u00e1 qu\u00e9 im\u00e1genes se utilizaron.')
    add_paragraph_justified(doc,
        'Hasta la cancelaci\u00f3n total del precio, el c\u00f3digo fuente permanece bajo propiedad del Desarrollador.')

    # ── CL\u00c1USULA 7 ──
    add_heading_styled(doc, 'Cl\u00e1usula S\u00e9ptima \u2014 Soporte post-entrega', level=2)
    add_paragraph_justified(doc, 'Per\u00edodo de soporte incluido seg\u00fan plan:')
    add_bullet(doc, 'Plan Starter: no incluye soporte post-entrega.')
    add_bullet(doc, 'Plan Profesional: 3 (tres) meses desde la entrega.')
    add_bullet(doc, 'Plan Premium: 6 (seis) meses desde la entrega.')
    add_paragraph_justified(doc, 'El soporte incluido cubre:', bold=True)
    add_bullet(doc, 'Correcci\u00f3n de errores t\u00e9cnicos (bugs) del sitio entregado.')
    add_bullet(doc, 'Cambios menores de contenido: textos, precios, fotos e informaci\u00f3n de contacto.')
    add_bullet(doc, 'Asistencia t\u00e9cnica por consultas relacionadas al sitio.')
    add_paragraph_justified(doc, 'El soporte NO incluye:', bold=True)
    add_bullet(doc, 'Desarrollo de funcionalidades nuevas.')
    add_bullet(doc, 'Redise\u00f1o parcial o total del sitio.')
    add_bullet(doc, 'Cambios estructurales (agregar secciones, cambiar arquitectura).')
    add_bullet(doc, 'Gesti\u00f3n de campa\u00f1as de marketing o redes sociales.')
    add_paragraph_justified(doc,
        'Finalizado el per\u00edodo de soporte, el Cliente puede contratar mantenimiento mensual '
        '(referencia: desde $15.000/mes al momento de la firma).')

    # ── CL\u00c1USULA 8 ──
    add_heading_styled(doc, 'Cl\u00e1usula Octava \u2014 Hosting y dominio', level=2)
    add_paragraph_justified(doc,
        'El sitio se publicar\u00e1 en una plataforma de hosting profesional sin costo mensual para '
        'sitios est\u00e1ticos (Vercel, Netlify, GitHub Pages o similar).')
    add_paragraph_justified(doc,
        'Si el proyecto requiere funcionalidades con servidor, el Desarrollador informar\u00e1 el costo '
        'mensual estimado antes de comenzar. Dicho costo ser\u00e1 a cargo del Cliente.')
    add_paragraph_justified(doc,
        'El registro del dominio (primer a\u00f1o) est\u00e1 incluido. La renovaci\u00f3n anual posterior es '
        'responsabilidad del Cliente. El dominio se registra a nombre del Cliente.')

    # ── CL\u00c1USULA 9 ──
    add_heading_styled(doc, 'Cl\u00e1usula Novena \u2014 Prototipo gratuito', level=2)
    add_paragraph_justified(doc,
        'El Desarrollador ofrece un prototipo visual sin costo ni compromiso. Si el Cliente decide '
        'no continuar, no tiene obligaci\u00f3n de pago. El prototipo es un dise\u00f1o navegable que muestra '
        'la apariencia del sitio; no constituye el producto terminado. El prototipo y su c\u00f3digo '
        'son propiedad del Desarrollador hasta la contrataci\u00f3n formal.')

    # ── CL\u00c1USULA 10 ──
    add_heading_styled(doc, 'Cl\u00e1usula D\u00e9cima \u2014 Rescisi\u00f3n', level=2)
    add_paragraph_justified(doc,
        'Cualquiera de las partes puede rescindir el contrato mediante notificaci\u00f3n escrita '
        '(email con acuse de recibo) con 7 d\u00edas corridos de preaviso.')
    add_paragraph_justified(doc, 'En caso de rescisi\u00f3n por el Cliente:')
    add_bullet(doc, 'Antes de la aprobaci\u00f3n del prototipo: sin costo.')
    add_bullet(doc, 'Despu\u00e9s del prototipo: abona el trabajo realizado, m\u00ednimo 50% del precio total.')
    add_paragraph_justified(doc,
        'En caso de rescisi\u00f3n por el Desarrollador: reintegrar\u00e1 los montos por trabajo no realizado '
        'dentro de los 15 d\u00edas h\u00e1biles.')
    add_paragraph_justified(doc,
        'El contrato se rescinde de pleno derecho si el Cliente no provee contenido dentro de los '
        '30 d\u00edas corridos desde la firma, salvo acuerdo en contrario.')

    # ── CL\u00c1USULA 11 ──
    add_heading_styled(doc, 'Cl\u00e1usula Und\u00e9cima \u2014 Confidencialidad', level=2)
    add_paragraph_justified(doc,
        'Ambas partes se comprometen a mantener confidencial toda informaci\u00f3n del otro que no sea '
        'de conocimiento p\u00fablico, incluyendo datos comerciales, estrategias, informaci\u00f3n de clientes '
        'y documentaci\u00f3n intercambiada durante la ejecuci\u00f3n del contrato.')

    # ── CL\u00c1USULA 12 ──
    add_heading_styled(doc, 'Cl\u00e1usula Duod\u00e9cima \u2014 Responsabilidad', level=2)
    add_paragraph_justified(doc,
        'El Desarrollador garantiza que el sitio funcionar\u00e1 correctamente seg\u00fan las especificaciones '
        'al momento de la entrega.')
    add_paragraph_justified(doc, 'El Desarrollador no es responsable por:')
    add_bullet(doc, 'Contenido provisto por el Cliente que infrinja derechos de terceros.')
    add_bullet(doc, 'Interrupciones causadas por la plataforma de hosting de terceros.')
    add_bullet(doc, 'Resultados comerciales del sitio (ventas, posicionamiento, tr\u00e1fico).')
    add_bullet(doc, 'Cambios realizados por el Cliente o terceros en el c\u00f3digo post-entrega.')
    add_paragraph_justified(doc,
        'La responsabilidad total del Desarrollador se limita al monto efectivamente percibido.')

    # ── CL\u00c1USULA 13 ──
    add_heading_styled(doc, 'Cl\u00e1usula Decimotercera \u2014 Resoluci\u00f3n de conflictos', level=2)
    add_paragraph_justified(doc,
        'Para cualquier controversia, las partes se someten a la jurisdicci\u00f3n de los Tribunales '
        'Ordinarios de la ciudad de Mendoza, Provincia de Mendoza, Rep\u00fablica Argentina, renunciando '
        'a cualquier otro fuero o jurisdicci\u00f3n.')

    # ── CL\u00c1USULA 14 ──
    add_heading_styled(doc, 'Cl\u00e1usula Decimocuarta \u2014 Disposiciones generales', level=2)
    add_paragraph_justified(doc,
        'El presente contrato constituye el acuerdo completo entre las partes y reemplaza cualquier '
        'acuerdo previo. Cualquier modificaci\u00f3n debe constar por escrito. La nulidad de alguna '
        'cl\u00e1usula no afecta las restantes. Se rige por las leyes de la Rep\u00fablica Argentina.')

    # ── Signatures ──
    doc.add_paragraph()
    add_paragraph_justified(doc,
        'En prueba de conformidad, se firman dos ejemplares de un mismo tenor y a un solo efecto.')
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(7.5)
    table.columns[1].width = Cm(7.5)

    table.cell(0, 0).text = 'EL DESARROLLADOR'
    table.cell(0, 1).text = 'EL CLIENTE'
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''
    table.cell(2, 0).text = '_________________________'
    table.cell(2, 1).text = '_________________________'
    table.cell(3, 0).text = 'Juan Ignacio Barranco\nDNI: _________________'
    table.cell(3, 1).text = 'Nombre: _______________\nDNI/CUIT: _____________'

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = 'Calibri'
            if row == table.rows[0]:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    # ── Page break -> Anexo A ──
    doc.add_page_break()

    title_a = doc.add_heading('ANEXO A \u2014 Alcance del proyecto', level=1)
    for run in title_a.runs:
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    add_blank_field(doc, 'Nombre del negocio', 40)
    add_blank_field(doc, 'Rubro', 40)
    doc.add_paragraph()
    add_paragraph_justified(doc, 'Plan/servicio contratado:')
    add_checkbox_line(doc, 'Starter ($79.200 \u2014 lanzamiento)')
    add_checkbox_line(doc, 'Profesional ($224.000 \u2014 lanzamiento)')
    add_checkbox_line(doc, 'Premium ($384.000 \u2014 lanzamiento)')
    add_checkbox_line(doc, 'Analytics Setup ($44.000)')
    add_checkbox_line(doc, 'Data Pro ($152.000)')
    add_checkbox_line(doc, 'Data Enterprise (a medida)')
    doc.add_paragraph()
    add_paragraph_justified(doc, 'Descripci\u00f3n del sitio/servicio:')
    for _ in range(3):
        add_paragraph_justified(doc, '_' * 80)
    doc.add_paragraph()
    add_paragraph_justified(doc, 'Secciones incluidas:', bold=True)
    for s in [
        'P\u00e1gina principal (hero, presentaci\u00f3n)',
        'Sobre nosotros / Historia',
        'Servicios / Productos / Men\u00fa / Carta',
        'Galer\u00eda de fotos',
        'Sistema de reservas / turnos online',
        'Cat\u00e1logo con integraci\u00f3n de pagos (Mercado Pago)',
        'Formulario de contacto',
        'Testimonios / Rese\u00f1as',
        'Ubicaci\u00f3n y horarios',
        'Versi\u00f3n en ingl\u00e9s',
    ]:
        add_checkbox_line(doc, s)
    add_checkbox_line(doc, 'Otra: ________________________________________')

    doc.add_paragraph()
    add_paragraph_justified(doc, 'Servicios de datos:', bold=True)
    add_checkbox_line(doc, 'Google Analytics 4 con eventos personalizados')
    add_checkbox_line(doc, 'Dashboard de m\u00e9tricas clave')
    add_checkbox_line(doc, 'Integraci\u00f3n de m\u00faltiples fuentes de datos')
    add_checkbox_line(doc, 'Automatizaci\u00f3n de reportes')
    add_checkbox_line(doc, 'Pipelines de datos personalizados')
    add_checkbox_line(doc, 'Data warehouse / data lake')
    add_checkbox_line(doc, 'Modelos predictivos')
    add_checkbox_line(doc, 'Otro: ________________________________________')

    doc.add_paragraph()
    add_paragraph_justified(doc, 'Funcionalidades especiales:')
    for _ in range(2):
        add_paragraph_justified(doc, '_' * 80)

    doc.add_paragraph()
    add_blank_field(doc, 'Dominio deseado', 35)

    doc.add_paragraph()
    add_paragraph_justified(doc, 'Contenido a cargo del Cliente:', bold=True)
    for item in [
        'Textos y descripciones',
        'Fotograf\u00edas propias',
        'Logotipo',
        'Informaci\u00f3n de contacto y horarios',
        'Lista de productos/servicios con precios',
    ]:
        add_checkbox_line(doc, item)

    doc.add_paragraph()
    add_blank_field(doc, 'Fecha l\u00edmite de entrega de contenido', 20)

    doc.add_paragraph()
    add_paragraph_justified(doc, 'Notas adicionales:')
    for _ in range(3):
        add_paragraph_justified(doc, '_' * 80)

    # ── Save ──
    path = os.path.join(OUT_DIR, 'contrato-servicios.docx')
    doc.save(path)
    print(f'DOCX saved: {path}')
    return path


# ── build PDF ────────────────────────────────────────────

def build_pdf():
    from fpdf import FPDF

    FONT_DIR = 'C:/Windows/Fonts/'

    # Line heights
    LH_BODY = 6
    LH_HEADING = 8
    LH_BULLET = 6
    LH_CHECKBOX = 6

    class ContratoPDF(FPDF):
        def __init__(self):
            super().__init__()
            self.add_font('Cal', '', FONT_DIR + 'calibri.ttf', uni=True)
            self.add_font('Cal', 'B', FONT_DIR + 'calibrib.ttf', uni=True)
            self.add_font('Cal', 'I', FONT_DIR + 'calibrii.ttf', uni=True)
            self.add_font('Cal', 'BI', FONT_DIR + 'calibriz.ttf', uni=True)

        def header(self):
            if self.page_no() == 1:
                return
            self.set_font('Cal', 'I', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 8, 'Contrato de Locaci\u00f3n de Servicios \u2014 Juan Barranco', align='C')
            self.ln(10)

        def footer(self):
            self.set_y(-15)
            self.set_font('Cal', 'I', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f'P\u00e1gina {self.page_no()}/{self.pages_count}', align='C')

        def remaining_space(self):
            """Return mm of usable space left on current page."""
            return self.h - self.get_y() - 25  # 25mm bottom margin

        def ensure_space(self, mm_needed):
            """Add a new page if not enough space remains."""
            if self.remaining_space() < mm_needed:
                self.add_page()

        def section_title(self, text):
            self.ensure_space(20)  # need at least 20mm for heading + some body
            self.set_font('Cal', 'B', 12)
            self.set_text_color(30, 58, 95)
            self.multi_cell(0, LH_HEADING, text)
            self.ln(2)

        def body_text(self, text):
            self.set_font('Cal', '', 10)
            self.set_text_color(40, 40, 40)
            self.multi_cell(0, LH_BODY, text)
            self.ln(1)

        def bold_text(self, text):
            self.set_font('Cal', 'B', 10)
            self.set_text_color(40, 40, 40)
            self.multi_cell(0, LH_BODY, text)
            self.ln(1)

        def bullet(self, text):
            self.set_font('Cal', '', 10)
            self.set_text_color(40, 40, 40)
            x = self.get_x()
            self.cell(6, LH_BULLET, '\u2022')
            self.multi_cell(0, LH_BULLET, text)
            self.ln(0.5)

        def checkbox(self, text):
            self.set_font('Cal', '', 10)
            self.set_text_color(40, 40, 40)
            self.cell(6, LH_CHECKBOX, '\u2610')
            self.multi_cell(0, LH_CHECKBOX, f'  {text}')
            self.ln(0.5)

        def field_line(self, label, width=120):
            self.set_font('Cal', 'B', 10)
            self.set_text_color(40, 40, 40)
            self.cell(40, LH_BODY, f'{label}: ')
            self.set_font('Cal', '', 10)
            self.set_draw_color(180, 180, 180)
            x = self.get_x()
            y = self.get_y() + 5
            self.line(x, y, x + width, y)
            self.ln(8)

    pdf = ContratoPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Title
    pdf.set_font('Cal', 'B', 20)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(0, 12, 'CONTRATO DE LOCACI\u00d3N DE SERVICIOS', align='C')
    pdf.ln(12)
    pdf.set_font('Cal', 'B', 14)
    pdf.cell(0, 10, 'DE DESARROLLO WEB Y DATOS', align='C')
    pdf.ln(10)
    # Blue line
    pdf.set_draw_color(59, 130, 246)
    pdf.set_line_width(0.5)
    pdf.line(30, pdf.get_y(), 180, pdf.get_y())
    pdf.ln(6)

    # Parties
    pdf.section_title('Partes')
    pdf.body_text(
        'EL PRESTADOR: Juan Ignacio Barranco, DNI N.\u00b0 ____________, '
        'con domicilio en Maip\u00fa, Mendoza, correo electr\u00f3nico jbarranco.sistemas@gmail.com, '
        'tel\u00e9fono +54 9 261 642-3730, en adelante \u201cEL DESARROLLADOR\u201d.')
    pdf.body_text(
        'EL COMITENTE: ____________________________________________, '
        'DNI/CUIT N.\u00b0 ____________, con domicilio comercial en '
        '____________________________________________, correo electr\u00f3nico '
        '________________________, tel\u00e9fono ________________________, en adelante \u201cEL CLIENTE\u201d.')
    pdf.body_text(
        'En la ciudad de Mendoza, a los ______ d\u00edas del mes de ______________ de 20____, '
        'las partes acuerdan celebrar el presente contrato de locaci\u00f3n de servicios conforme '
        'a los art\u00edculos 1251 a 1261 del C\u00f3digo Civil y Comercial de la Naci\u00f3n.')

    # Clause 1
    pdf.section_title('Cl\u00e1usula Primera \u2014 Objeto')
    pdf.body_text(
        'El Desarrollador se compromete a dise\u00f1ar, desarrollar y entregar al Cliente un sitio '
        'web profesional y/o servicios de datos seg\u00fan las especificaciones del Anexo A \u2014 '
        'Alcance del proyecto.')
    pdf.body_text('Plan/servicio contratado:')
    pdf.checkbox('Starter ($79.200 \u2014 lanzamiento)')
    pdf.checkbox('Profesional ($224.000 \u2014 lanzamiento)')
    pdf.checkbox('Premium ($384.000 \u2014 lanzamiento)')
    pdf.checkbox('Analytics Setup ($44.000)')
    pdf.checkbox('Data Pro ($152.000)')
    pdf.checkbox('Data Enterprise (a medida)')

    # Clause 2
    pdf.section_title('Cl\u00e1usula Segunda \u2014 Alcance y exclusiones')
    pdf.bold_text('2.1. Incluye:')
    pdf.bold_text('Plan Starter ($79.200 \u2014 lanzamiento):')
    for item in [
        'Sitio web profesional a medida (hasta 5 secciones)',
        'Dise\u00f1o responsive (celular, PC, tablet)',
        'Formulario de contacto funcional + bot\u00f3n de WhatsApp',
        'Google My Business optimizado + SEO b\u00e1sico',
        'Google Analytics 4 configurado',
        'Dominio .com.ar primer a\u00f1o + hosting sin costo mensual',
    ]:
        pdf.bullet(item)

    pdf.bold_text('Plan Profesional ($224.000 \u2014 lanzamiento) \u2014 todo lo anterior m\u00e1s:')
    for item in [
        'Sitio de 5+ secciones',
        'Sistema de reservas o turnos online integrado',
        'SEO avanzado + Google Analytics completo',
        'Capacitaci\u00f3n para gestionar contenido',
    ]:
        pdf.bullet(item)

    pdf.bold_text('Plan Premium ($384.000 \u2014 lanzamiento) \u2014 todo lo anterior m\u00e1s:')
    for item in [
        'Cat\u00e1logo con integraci\u00f3n de Mercado Pago',
        'Versi\u00f3n biling\u00fce (espa\u00f1ol e ingl\u00e9s)',
        'Google Analytics avanzado con eventos personalizados',
        'Formularios inteligentes con notificaciones autom\u00e1ticas',
    ]:
        pdf.bullet(item)

    # Data services in PDF
    pdf.bold_text('Servicios de datos:')

    pdf.bold_text('Analytics Setup ($44.000):')
    for item in [
        'Configuraci\u00f3n de Google Analytics 4 con eventos personalizados',
        'Dashboard b\u00e1sico de m\u00e9tricas clave',
        'Integraci\u00f3n con Google Search Console',
        'Informe inicial de rendimiento',
    ]:
        pdf.bullet(item)

    pdf.bold_text('Data Pro ($152.000):')
    for item in [
        'Todo lo de Analytics Setup',
        'Dashboard avanzado con m\u00e9tricas de negocio',
        'Integraci\u00f3n de m\u00faltiples fuentes de datos',
        'Automatizaci\u00f3n de reportes peri\u00f3dicos',
        'Capacitaci\u00f3n para interpretar datos',
    ]:
        pdf.bullet(item)

    pdf.bold_text('Data Enterprise (desde $304.000 \u2014 a medida):')
    for item in [
        'Todo lo de Data Pro',
        'Pipelines de datos personalizados',
        'Data warehouse / data lake seg\u00fan necesidad',
        'Modelos predictivos y an\u00e1lisis avanzado',
        'Soporte continuo y mantenimiento de infraestructura de datos',
    ]:
        pdf.bullet(item)

    pdf.bold_text('2.2. No incluye:')
    for item in [
        'Creaci\u00f3n de contenido (textos, fotos, logo) \u2014 el Cliente los provee',
        'Gesti\u00f3n de redes sociales ni campa\u00f1as de publicidad',
        'Funcionalidades no especificadas en el Anexo A',
        'Aplicaciones m\u00f3viles nativas',
    ]:
        pdf.bullet(item)

    # Clause 3
    pdf.section_title('Cl\u00e1usula Tercera \u2014 Precio y forma de pago')
    pdf.field_line('Precio total')
    pdf.body_text('Forma de pago:')
    pdf.checkbox('50% al aprobar prototipo, 50% al entregar')
    pdf.checkbox('3 cuotas de $__________ (primera al aprobar prototipo)')
    pdf.checkbox('Otra: ________________________________________')
    pdf.body_text(
        '3.1. Medios de pago aceptados: transferencia bancaria, efectivo, Mercado Pago.')
    pdf.body_text(
        '3.2. Stablecoins aceptadas: USDT, USDC, DAI (red Tron o Polygon).')
    pdf.body_text(
        '3.3. Los pagos en criptomonedas se liquidan al tipo de cambio del momento de acreditaci\u00f3n.')
    pdf.body_text(
        '3.4. El Desarrollador no emite factura fiscal al momento de la firma. '
        'Si en el futuro se inscribiera como contribuyente, podr\u00e1 emitir el comprobante correspondiente.')
    pdf.body_text(
        '3.5. La falta de pago habilita a suspender el trabajo y retener archivos.')
    pdf.body_text(
        '3.6. Los precios indicados corresponden a la promoci\u00f3n de lanzamiento (20% de descuento '
        'sobre precios de lista), v\u00e1lida para los primeros 10 clientes. Precios de lista: '
        'Starter $99.000, Profesional $280.000, Premium $480.000.')

    # Clause 4
    pdf.section_title('Cl\u00e1usula Cuarta \u2014 Plazos')
    pdf.body_text('Prototipo visual: 48 horas h\u00e1biles desde recepci\u00f3n del contenido completo.')
    pdf.body_text('Sitio terminado (desde aprobaci\u00f3n del prototipo):')
    pdf.bullet('Starter: 7 d\u00edas h\u00e1biles')
    pdf.bullet('Profesional: 14 d\u00edas h\u00e1biles')
    pdf.bullet('Premium: 21 d\u00edas h\u00e1biles')
    pdf.body_text(
        'Los plazos se suspenden por demoras del Cliente. Se incluyen 2 rondas de revisi\u00f3n; '
        'cambios de estructura o alcance son trabajo adicional.')

    # Clause 5
    pdf.section_title('Cl\u00e1usula Quinta \u2014 Obligaciones del Cliente')
    pdf.bullet('Proveer contenido completo dentro de 5 d\u00edas h\u00e1biles de la firma.')
    pdf.bullet('Designar persona de contacto con capacidad de decisi\u00f3n.')
    pdf.bullet('Responder consultas en m\u00e1ximo 3 d\u00edas h\u00e1biles.')
    pdf.bullet('Facilitar acceso a cuentas necesarias.')
    pdf.bullet('Abonar el precio en los plazos convenidos.')

    # Clause 6
    pdf.section_title('Cl\u00e1usula Sexta \u2014 Propiedad intelectual')
    pdf.body_text(
        'Una vez cancelado el 100% del precio, todo el c\u00f3digo, dise\u00f1o y archivos pasan a ser '
        'propiedad exclusiva del Cliente, sin restricciones. El Desarrollador puede usar capturas '
        'en su portfolio, salvo prohibici\u00f3n escrita. Las im\u00e1genes stock se rigen por sus licencias. '
        'Hasta el pago total, el c\u00f3digo es propiedad del Desarrollador.')

    # Clause 7
    pdf.section_title('Cl\u00e1usula S\u00e9ptima \u2014 Soporte post-entrega')
    pdf.bullet('Starter: sin soporte post-entrega.')
    pdf.bullet('Profesional: 3 meses (bugs + cambios menores de contenido).')
    pdf.bullet('Premium: 6 meses (bugs + cambios menores de contenido).')
    pdf.body_text(
        'No incluye: funcionalidades nuevas, redise\u00f1o, cambios estructurales, marketing. '
        'Mantenimiento posterior: desde $15.000/mes.')

    # Clause 8
    pdf.section_title('Cl\u00e1usula Octava \u2014 Hosting y dominio')
    pdf.body_text(
        'Hosting sin costo para sitios est\u00e1ticos. Si se requiere servidor, el costo se informa '
        'antes de comenzar y es a cargo del Cliente. El dominio (primer a\u00f1o incluido) se registra '
        'a nombre del Cliente. Renovaci\u00f3n anual posterior a cargo del Cliente.')

    # Clause 9
    pdf.section_title('Cl\u00e1usula Novena \u2014 Prototipo gratuito')
    pdf.body_text(
        'Prototipo visual sin costo ni compromiso. Si el Cliente no contin\u00faa, no hay pago. '
        'Es un dise\u00f1o navegable, no el producto terminado. Es propiedad del Desarrollador '
        'hasta la contrataci\u00f3n formal.')

    # Clause 10
    pdf.section_title('Cl\u00e1usula D\u00e9cima \u2014 Rescisi\u00f3n')
    pdf.body_text('Rescisi\u00f3n con 7 d\u00edas corridos de preaviso por escrito (email con acuse).')
    pdf.bullet('Por el Cliente antes del prototipo: sin costo.')
    pdf.bullet('Por el Cliente despu\u00e9s del prototipo: m\u00ednimo 50% del precio total.')
    pdf.bullet('Por el Desarrollador: reintegro de montos por trabajo no realizado en 15 d\u00edas h\u00e1biles.')
    pdf.body_text('Rescisi\u00f3n autom\u00e1tica si no hay contenido en 30 d\u00edas corridos.')

    # Clause 11-14
    pdf.section_title('Cl\u00e1usula Und\u00e9cima \u2014 Confidencialidad')
    pdf.body_text('Ambas partes mantienen confidencial toda informaci\u00f3n no p\u00fablica del otro.')

    pdf.section_title('Cl\u00e1usula Duod\u00e9cima \u2014 Responsabilidad')
    pdf.body_text(
        'El Desarrollador garantiza funcionamiento seg\u00fan especificaciones. No responde por: '
        'contenido del Cliente, interrupciones de hosting, resultados comerciales, ni cambios '
        'post-entrega. Responsabilidad limitada al monto percibido.')

    pdf.section_title('Cl\u00e1usula Decimotercera \u2014 Jurisdicci\u00f3n')
    pdf.body_text('Tribunales Ordinarios de Mendoza, Argentina.')

    pdf.section_title('Cl\u00e1usula Decimocuarta \u2014 Disposiciones generales')
    pdf.body_text(
        'Acuerdo completo. Modificaciones por escrito. Nulidad parcial no afecta el resto. '
        'Se rige por leyes de la Rep\u00fablica Argentina.')

    # Signatures
    pdf.ensure_space(60)  # Need at least 60mm for signature block
    pdf.body_text('En prueba de conformidad, se firman dos ejemplares de un mismo tenor.')
    pdf.ln(10)

    y_sig = pdf.get_y()
    pdf.set_font('Cal', 'B', 10)
    pdf.set_text_color(40, 40, 40)

    pdf.set_xy(25, y_sig)
    pdf.cell(70, 6, 'EL DESARROLLADOR', align='C')
    pdf.set_xy(115, y_sig)
    pdf.cell(70, 6, 'EL CLIENTE', align='C')

    pdf.set_xy(25, y_sig + 20)
    pdf.cell(70, 6, '_________________________', align='C')
    pdf.set_xy(115, y_sig + 20)
    pdf.cell(70, 6, '_________________________', align='C')

    pdf.set_font('Cal', '', 9)
    pdf.set_xy(25, y_sig + 27)
    pdf.cell(70, 5, 'Juan Ignacio Barranco', align='C')
    pdf.set_xy(115, y_sig + 27)
    pdf.cell(70, 5, 'Nombre: _______________', align='C')

    pdf.set_xy(25, y_sig + 32)
    pdf.cell(70, 5, 'DNI: _________________', align='C')
    pdf.set_xy(115, y_sig + 32)
    pdf.cell(70, 5, 'DNI/CUIT: _____________', align='C')

    # ── Anexo A ──
    pdf.add_page()
    pdf.set_font('Cal', 'B', 16)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(0, 12, 'ANEXO A \u2014 Alcance del proyecto', align='C')
    pdf.ln(12)

    pdf.field_line('Nombre del negocio')
    pdf.field_line('Rubro')
    pdf.ln(2)
    pdf.body_text('Plan/servicio contratado:')
    pdf.checkbox('Starter ($79.200 \u2014 lanzamiento)')
    pdf.checkbox('Profesional ($224.000 \u2014 lanzamiento)')
    pdf.checkbox('Premium ($384.000 \u2014 lanzamiento)')
    pdf.checkbox('Analytics Setup ($44.000)')
    pdf.checkbox('Data Pro ($152.000)')
    pdf.checkbox('Data Enterprise (a medida)')
    pdf.ln(2)
    pdf.body_text('Descripci\u00f3n del sitio/servicio:')
    for _ in range(3):
        pdf.set_draw_color(180, 180, 180)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(7)
    pdf.ln(2)
    pdf.bold_text('Secciones incluidas:')
    for s in [
        'P\u00e1gina principal', 'Sobre nosotros / Historia',
        'Servicios / Productos / Men\u00fa', 'Galer\u00eda de fotos',
        'Reservas / Turnos online', 'Cat\u00e1logo con Mercado Pago',
        'Formulario de contacto', 'Testimonios', 'Ubicaci\u00f3n y horarios',
        'Versi\u00f3n en ingl\u00e9s', 'Otra: ____________________',
    ]:
        pdf.checkbox(s)

    pdf.ln(2)
    pdf.bold_text('Servicios de datos:')
    for s in [
        'Google Analytics 4 con eventos personalizados',
        'Dashboard de m\u00e9tricas clave',
        'Integraci\u00f3n de m\u00faltiples fuentes de datos',
        'Automatizaci\u00f3n de reportes',
        'Pipelines de datos personalizados',
        'Data warehouse / data lake',
        'Modelos predictivos',
        'Otro: ____________________',
    ]:
        pdf.checkbox(s)

    pdf.ln(2)
    pdf.field_line('Dominio deseado')
    pdf.ln(2)
    pdf.bold_text('Contenido a cargo del Cliente:')
    for item in ['Textos y descripciones', 'Fotograf\u00edas', 'Logotipo',
                 'Contacto y horarios', 'Lista de productos/precios']:
        pdf.checkbox(item)
    pdf.ln(2)
    pdf.field_line('Fecha l\u00edmite contenido')
    pdf.ln(2)
    pdf.body_text('Notas adicionales:')
    for _ in range(3):
        pdf.set_draw_color(180, 180, 180)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(7)

    path = os.path.join(OUT_DIR, 'contrato-servicios.pdf')
    pdf.output(path)
    print(f'PDF saved: {path}')
    return path


if __name__ == '__main__':
    build_docx()
    build_pdf()
    print('Done!')
